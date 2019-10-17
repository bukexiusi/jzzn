
import re, os, copy, time
from utils import DateHelper, WordHelper
'''辅助函数'''
import win32com.client


def processExist(processName):
    try:
        WMI = win32com.client.GetObject('winmgmts:')
        processCodeCov = WMI.ExecQuery('select * from Win32_Process where Name="%s"' % processName)

        if len(processCodeCov) > 0:
            print(processName + " exist")
            return True
        else:
            print(processName + " is not exist")
            return False
    except Exception as e:
        print(processName + "error : ", e)
        raise e


def killProcess(processName):
    if processExist(processName):
        os.system('taskkill /f /im %s' % processName)


def valueFormatDocument(mdata_key, mdata, word, input_path, extra_picture_dir_path):
    mdata_key_temp = re.findall("([\u4e00-\u9fa5\da-zA-Z\-\_]*?)\(zn\:", mdata_key)
    if len(mdata_key_temp) > 0:
        mdata_key_true = mdata_key_temp[0]
        var = mdata.get(mdata_key_true)
        if not var:
            return var
        if "(zn:date " in mdata_key:
            try:
                time.strptime(var, "%Y-%m-%d %H:%M:%S")
            except Exception:
                try:
                    time.strptime(var, "%Y-%m-%d")
                    var = var + " 00:00:00"
                except Exception:
                    try:
                        time.strptime(var, "%y-%m-%d")
                        var = "20" + var + " 00:00:00"
                    except Exception:
                        pass
            mdata_key_format = re.findall("\(zn:date (.*?)\)", mdata_key)
            if len(mdata_key_format) < 1:
                raise Exception(input_path + "中" + mdata_key + "处可能缺少必要空格")
            mdata_key_format = mdata_key_format[0]
            if mdata_key_format == "中文日期零":
                var = DateHelper.chinese_date(var)
            elif mdata_key_format == "中文日期〇":
                var = DateHelper.chinese_date2(var)
            else:
                time_str = var
                try:
                    if time_str:
                        time_array = time.strptime(time_str, "%Y-%m-%d %H:%M:%S")
                        var = time.strftime(mdata_key_format, time_array).format(y='年', m='月',
                                                                                 d='日')
                    else:
                        var = ""
                except ValueError as e:
                    time_array = time.strptime(time_str, "%Y-%m-%d")
                    var = time.strftime(mdata_key_format, time_array).format(y='年', m='月',
                                                                             d='日')
                except Exception as e:
                    var = mdata.get(mdata_key)
            # 去掉月份前面0，日期前面0（南安需求），同时也会出去时分秒前面零
            re_var = re.findall("\d+", var)
            if len(re_var) > 0:
                var = re.sub("\d+", "{zn}", var)
                for rv in re_var:
                    if rv.startswith("0"):
                        rv = rv[1:]
                    var = re.sub("{zn}", rv, var, 1)

        if "(zn:format " in mdata_key:
            if not ("zn:endformat" in mdata_key):
                raise Exception(input_path + "中" + mdata_key + "处可能缺少zn:endformat")
            mdata_key_format = re.findall("\(zn:format (.*) zn:endformat\)", mdata_key)
            if len(mdata_key_format) < 1:
                raise Exception(input_path + "中" + mdata_key + "缺少必要空格")
            mdata_key_format = mdata_key_format[0]
            mdata_key_format = mdata_key_format.replace('“', '"').replace('”', '"').replace("‘", "'").replace("’",
                                                                                                              "'")
            var = eval(mdata_key_format)

        if "(zn:func " in mdata_key:
            ''' mdata_key ->  '''
            mdata_key_format = re.findall("\(zn:func (.*)\)", mdata_key)
            if len(mdata_key_format) < 1:
                raise Exception(input_path + "中" + mdata_key + "缺少必要空格")
            var = eval(mdata_key_format[0] + '("' + var + '")')

        if "(zn:picture)" in mdata_key:
            if extra_picture_dir_path:
                if extra_picture_dir_path.endswith("/"):
                    picture_path = extra_picture_dir_path + var + ".bmp"
                else:
                    picture_path = extra_picture_dir_path + "/" + var + ".bmp"
                word.replace_picture("[" + mdata_key + "]", picture_path)
            return mdata_key

        if "(zn:name)" in mdata_key:  # 不做格式上的处理
            pass
    else:
        var = mdata.get(mdata_key)
        if var is None:
            var = mdata_key
    return var
'''辅助函数'''


'''begin：word操作'''
from docx import Document
from win32com.client import Dispatch
import sys

class ZNWord():

    def __init__(self, path=None, delDocx=False):
        self.word = None
        self.fullPath = None
        self.dealtPath = None
        if path:
            self.open(path)

    '''将doc转存为docx'''
    def docToDocx(self, docPath, docxPath):
        word = Dispatch('word.Application')
        # word = win32com.client.gencache.EnsureDispatch('kwps.application')
        doc = word.Documents.Open(docPath)
        doc.SaveAs(docxPath, FileFormat=12)
        doc.Close(0)
        word.Quit()

    '''打开word'''
    def open(self, path):
        path = path.replace("\\", "/")
        self.fullPath = path
        try:
            if path.endswith(".docx"):
                self.word = Document(path)
            elif path.endswith(".doc"):
                docxPath = self.fullPath.replace(".doc", ".docx")
                self.docToDocx(self.fullPath, docxPath)
                self.dealtPath = docxPath
                self.word = Document(docxPath)
            else:
                raise Exception("".join([path, "不是doc类型或者docx类型文件"]))

        except ValueError:
            raise Exception("word文档不是docx结构（doc文档需要另存为docx，不能直接改后缀）")

    '''word内容'''
    def content(self):
        self.__checkWord()
        content = ""
        for p in self.word.paragraphs:
            content = content + p.text
        return content

    '''替换word内容'''
    def replace(self, oldText, newText):
        self.__checkWord()
        # 段落替换
        self.__replaceParagraphs(oldText, newText, self.word.paragraphs)
        # 表格替换
        for table in self.word.tables:
            for cell in table._cells:
                if oldText in cell.text:
                    self.__replaceParagraphs(oldText, newText, cell.paragraphs)

    def __replaceParagraphs(self, oldText, newText, paragraphs):
        for p in paragraphs:
            pText = p.text
            subTag = None
            replaceMappingArray = []
            if oldText in pText:
                oldTempText = oldText
                runIndex = 0
                for run in p.runs:
                    runText = run.text
                    if (not subTag) and oldTempText in runText:
                        text = runText.replace(oldTempText, newText)
                        run.text = text
                    else:
                        if not subTag:
                            firstCharacter = oldTempText[0]
                            try:
                                start = runText.index(firstCharacter)
                            except ValueError:
                                continue

                            matchStr = runText[start:]
                            while True:
                                if oldTempText.startswith(matchStr):
                                    subTag = "start"
                                    oldTempText = oldTempText.replace(matchStr, "", 1)
                                    replaceMapping = {
                                        "runIndex": runIndex,
                                        "replaceText": matchStr
                                    }
                                    replaceMappingArray.append(replaceMapping)
                                    break
                                else:
                                    start = start + 1
                                    if start == len(oldText):
                                        # 无匹配项
                                        break
                                    try:
                                        matchStr = matchStr[start:]
                                    except ValueError:
                                        break
                        else:
                            if (oldTempText == runText) or runText.startswith(oldTempText):
                                subTag = "end"
                                replaceMapping = {
                                    "runIndex": runIndex,
                                    "replaceText": oldTempText
                                }
                                replaceMappingArray.append(replaceMapping)
                            elif oldTempText.startswith(runText):
                                oldTempText = oldTempText.replace(runText, "", 1)
                                subTag = "middle"
                                replaceMapping = {
                                    "runIndex": runIndex,
                                    "replaceText": runText
                                }
                                replaceMappingArray.append(replaceMapping)
                    runIndex = runIndex + 1
                if subTag and subTag == "end":
                    for index, val in enumerate(replaceMappingArray):
                        runIndex = val.get("runIndex")
                        runText = p.runs[runIndex].text
                        if index == 0:
                            text = runText.replace(val.get("replaceText"), newText)
                        else:
                            text = runText.replace(val.get("replaceText"), "")
                        p.runs[runIndex].text = text

    '''保存'''
    def save(self, path=None):
        if path:
            '''另存为'''
            path = path.replace("\\", "/")
            dir = re.findall(r"(.*)/", path)[0]
            if not os.path.exists(dir):
                os.makedirs(dir)
            self.word.save(path)
        else:
            '''保存'''
            self.word.save(self.fullPath)

    '''退出'''
    def quit(self):
        if self.dealtPath:
            os.remove(self.dealtPath)

    def __checkWord(self):
        if not self.word:
            raise Exception("没有执行open方法打开word")

'''end：word操作'''


'''begin：excel操作'''
import xlrd


class ZNExcel():

    def __init__(self, path=None):
        self.excel = None
        self.sheetNames = []
        self.path = path
        if path:
            self.open(path)
        self.temp = {}
        self.columnMapping = {}

    '''打开'''
    def open(self, path):
        self.path = path
        self.excel = xlrd.open_workbook(path)
        self.sheetNames = self.excel.sheet_names()

    '''获取sheet对象'''
    def getSheet(self, arg=None):

        if not arg:
            sheetName = self.sheetNames[0]
            sheetIndex = 0
            return ZNExcelSheet(self.excel, sheetName, sheetIndex)

        if isinstance(arg, str):
            sheetName = arg
            if not (sheetName in self.sheetNames):
                raise Exception(self.path + "中不存在名为" + sheetName + "的sheet名称")
            sheetIndex = self.sheetNames.index(sheetName)
        elif isinstance(arg, int):
            sheetIndex = arg
            sheetName = self.sheetNames[sheetIndex]
        else:
            raise Exception("看到这个异常说明你适合测试工作")
        return ZNExcelSheet(self.excel, sheetName, sheetIndex)


class ZNExcelSheet():
    '''初始化'''
    def __init__(self, excel, sheetName, sheetIndex):
        self.excel = excel
        self.sheetName = sheetName
        self.sheetIndex = sheetIndex
        self.sheet = excel.sheet_by_index(sheetIndex)
        self.columnNames = self.sheet.row_values(0)

    def getRowByIndex(self, idx):
        '''
        :return: 数组类型
        '''
        return self.sheet.row_values(idx)

    def getDictByIndex(self, idx):
        '''
        :return: 字典类型
        '''
        dictResult = {}
        if idx == 0:
            return dictResult
        values = self.sheet.row_values(idx)
        for i in range(len(self.columnNames)):
            if isinstance(values[i], float):
                dictResult[self.columnNames[i]] = int(values[i])
            else:
                dictResult[self.columnNames[i]] = values[i]
        return dictResult

    def getDictByCondition(self, columnName, columnValue):
        if not (columnName in self.columnNames):
            return {}
        columnIdx = self.columnNames.index(columnName)
        columnValues = self.sheet.col_values(columnIdx)
        columnValues = sorted(enumerate(columnValues), key=lambda x: columnValue in x, reverse=True)
        vIdxArray = []
        i = 0
        while True:
            vIdx, vVal = columnValues[i]
            if vVal != columnValue:
                break
            vIdxArray.append(vIdx)
            i = i + 1
        if len(vIdxArray) == 0:
            return {}

        if len(vIdxArray) == 1:
            return self.getDictByIndex(vIdxArray[0])
        else:
            result = []
            for element in vIdxArray:
                result.append(self.getDictByIndex(element))
            return result

    def setDictByIndex(self, idx, dictObj):
        pass

'''end：excel操作'''

if __name__ == "__main__":
    # excel = ZNExcel(r"D:\chunsun\列表导出.xls")
    # '''excel所有sheet名'''
    # print(excel.sheetNames)
    # '''获取sheet对象'''
    # sheet = excel.getSheet()
    # '''所有列名'''
    # print(sheet.columnNames)
    # '''获取某行数据'''
    # print(sheet.getRowByIndex(1))
    # '''获取某行字典'''
    # print(sheet.getDictByIndex(1))
    # '''根据条件获取某行'''
    # print(sheet.getDictByCondition("案号", "(2019)闽0203执3469号"))
    # print(sheet.getDictByCondition("案号", "(2019)闽0203执3470号"))

    # 替换测试
    word = ZNWord("D:\chunsun\纳入失信决定书.docx")
    wordText = word.content()
    re_array2 = re.findall(r"\{\{([\s\S]*?)\}\}", wordText)
    oldText = "{{" + re_array2[0] + "}}"
    newText = "jzznjzznjzznjzznjzznjzznjzznjzznjzznjzznjzznjzznjzznjzznjzzn"
    word.replace(oldText, newText)
    word.replace("[案号]", "(2019)闽0203执123号")
    word.replace("[文书生成日期(zn:date 中文日期〇)]", "2019年4月26号")
    word.save("D:\chunsun\纳入失信决定书2.docx")

    # word = ZNWord("D:/chunsun/案件模板/执行裁定书(判决).docx")
    # word.replace("[案号]", "(2019)闽0203执123号")
    # word.sava("D:/chunsun/案件模板/执行裁定书(判决)2.docx")

    # word.docToDocx("D:/chunsun/案件模板/执行裁定书(判决).doc", "D:/chunsun/案件模板/执行裁定书(判决)3.docx")
    # word2 = ZNWord("D:/chunsun/案件模板/执行裁定书(判决)3.docx")

    # word = ZNWord("D:/chunsun/案件模板/执行裁定书(判决).doc")
    # word2 = ZNWord("D:/chunsun/案件模板/执行裁定书(判决).docx")