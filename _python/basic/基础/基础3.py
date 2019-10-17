import docx
script = "import re"
exec(script)
print(re.findall(".*", "123"))

path = "D:/生成路径/2019年4月24日/(2019)闽0203执210号/执行裁定书.docx"
path = re.sub("\[.*?\]", ".", path)
print(path)

# old_text = "[案号]"
# new_text = "我就去了"
# file = docx.Document("D:\chunsun\纳入失信决定书.docx")
# 遍历文件对象
# for f in file.paragraphs:
#     # 如果 旧字符串 在 某个段落 中
#     if old_text in f.text:
#         print("替换前:", f.text)
#         # 将段落存入 inline
#         inline = f.runs
#         # 遍历 段落 生成 i
#         for i in inline:
#             # 如果 旧字符串 在 i 中
#             if old_text in i.text:
#                 # 替换 i.text 内文本资源
#                 text = i.text.replace(old_text, new_text)
#                 i.text = text
#         print("替换后===>", f.text)
# file.save("D:\chunsun\纳入失信决定书1.doc")

document = docx.Document("D:\chunsun\纳入失信决定书.docx")
wordText = ""
for p in document.paragraphs:
    wordText = wordText + p.text
print(wordText)

re_array2 = re.findall(r"\{\{([\s\S]*?)\}\}", wordText)
oldText = "{{"+re_array2[0]+"}}"
newText = "jzzn"
print(re_array2)
    # pText = ""
    # for run in p.runs
document.save("D:\chunsun\纳入失信决定书2.docx")

try:
    aaaa = "abcacccccccccccccccccc"
    print(aaaa.index("a", 4))
except ValueError:
    print("1111")

def a():
    return "1"

print(str(a))